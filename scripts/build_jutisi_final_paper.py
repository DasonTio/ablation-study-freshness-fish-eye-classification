#!/usr/bin/env python3
"""Build the final JuTISI manuscript as paper.docx.

The manuscript is intentionally generated from code so tables, figures, margins,
and references stay reproducible from the repo result files.
"""
from __future__ import annotations

import json
import subprocess
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "paper.docx"
FONT = "Times New Roman"
TITLE_ID = "Statistik Okular Global Memvalidasi GAP Swin-Tiny untuk Klasifikasi Kesegaran Mata Ikan"
TITLE_EN = "Global Ocular Statistics Validate GAP Swin-Tiny for Fish-Eye Freshness Classification"
FINAL_MODEL_STATEMENT = (
    "Final model: GAP Swin-Tiny + 24-class head, trained end-to-end without CLAHE."
)
GLOBAL_SIGNAL_ROLE = (
    "The global-statistics classifier is diagnostic evidence, not the final deployment model."
)


def run_figure_builder() -> None:
    subprocess.run(["python", str(ROOT / "scripts" / "make_jutisi_figures.py")], check=True)
    subprocess.run([
        "python",
        str(ROOT / "scripts" / "make_split_xai_figures.py"),
        "--skip-if-exists",
    ], check=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 8) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold


def apply_document_style(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(4.3)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.43)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_paragraph(doc: Document, text: str = "", *, indent: bool = True, bold: bool = False,
                  italic: bool = False, size: int = 10, align=None) -> None:
    p = doc.add_paragraph()
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.55)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(24)
    run.bold = False


def add_center(doc: Document, text: str, *, size: int = 10, bold: bool = False,
               italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(10)
    run.bold = True
    run.font.small_caps = True


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(10)
    run.bold = True
    run.italic = True


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(8)


def add_figure(doc: Document, rel_path: str, caption: str, width_inches: float = 5.8) -> None:
    path = ROOT / rel_path
    if not path.exists():
        add_center(doc, f"[Missing figure: {rel_path}]", size=8, italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    add_caption(doc, caption)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def load_results() -> dict:
    clahe = pd.read_csv(ROOT / "results" / "multiseed_summary.csv")
    second = json.loads((ROOT / "results" / "secondorder_merged" / "secondorder_significance.json").read_text())
    ordinal = json.loads((ROOT / "results" / "stats_strengthening.json").read_text())
    global_summary = json.loads((ROOT / "results" / "global_signal" / "global_signal_summary.json").read_text())
    global_classifier = pd.read_csv(ROOT / "results" / "global_signal" / "global_signal_classifier_summary.csv")
    global_comparisons = pd.read_csv(ROOT / "results" / "global_signal" / "global_signal_classifier_comparisons.csv")
    return {
        "clahe": clahe,
        "second": second,
        "ordinal": ordinal,
        "global_summary": global_summary,
        "global_classifier": global_classifier,
        "global_comparisons": global_comparisons,
    }


def fmt_pm(mean: float, std: float) -> str:
    return f"{mean:.2f} +/- {std:.2f}"


def build_tables(results: dict) -> dict[str, list[list[str]]]:
    summary = {row["arm"]: row for row in results["second"]["summary"]}
    gap = summary["A_gap"]
    table_model = [
        [
            "Recommended classifier",
            "ImageNet-pretrained Swin-Tiny + GAP + 24-class species-freshness head",
            "Final model used for the main FFE recommendation",
        ],
        [
            "Input policy",
            "Raw RGB eye image; no CLAHE",
            "Preserves global ocular luminance/color statistics",
        ],
        [
            "Validation protocol",
            "64/16/20 split, five seeds",
            f'{gap["class_acc_mean"]:.2f} +/- {gap["class_acc_std"]:.2f}% 24-class accuracy',
        ],
        [
            "Prior-work context",
            "Higher than Hoang et al. 85.99% by 2.54 pp",
            "Indicative only because the exact split/protocol is not reproduced",
        ],
    ]

    global_summary = results["global_summary"]
    global_classifier = results["global_classifier"].set_index("feature_set")
    global_comparisons = results["global_comparisons"].set_index(["baseline", "candidate"])
    top_feature = global_summary["top_raw_center_features"][0]
    raw_center = global_classifier.loc["raw_center_global"]
    clahe_center = global_classifier.loc["clahe_center_global"]
    species_only = global_classifier.loc["species_only"]
    raw_vs_clahe = global_comparisons.loc[("clahe_center_global", "raw_center_global")]
    table0 = [
        [
            "Top central luminance statistic",
            f'{top_feature["feature"].replace("raw_center_", "")}; rho={top_feature["global_rho"]:.3f}, p<1e-60',
            f'Consistent in {top_feature["consistent_species"]}/{top_feature["species_n"]} species',
        ],
        [
            "Raw central global statistics",
            f'{raw_center["accuracy_mean"] * 100:.2f} +/- {raw_center["accuracy_std"] * 100:.2f}% freshness acc.; QWK {raw_center["qwk_mean"]:.3f}',
            "Direct global signal carries usable freshness information",
        ],
        [
            "CLAHE central global statistics",
            f'{clahe_center["accuracy_mean"] * 100:.2f} +/- {clahe_center["accuracy_std"] * 100:.2f}% freshness acc.; QWK {clahe_center["qwk_mean"]:.3f}',
            f'Raw beats CLAHE by {raw_vs_clahe["accuracy_delta_mean"] * 100:.2f} pp, p={raw_vs_clahe["paired_p"]:.3f}',
        ],
        [
            "Species-only baseline",
            f'{species_only["accuracy_mean"] * 100:.2f} +/- {species_only["accuracy_std"] * 100:.2f}% freshness acc.; QWK {species_only["qwk_mean"]:.3f}',
            "Freshness signal is not explained by species labels alone",
        ],
    ]

    clahe = results["clahe"]
    backbone_order = [
        ("resnet50", "ResNet50"),
        ("efficientnetv2s", "EfficientNetV2-S"),
        ("convnext_small", "ConvNeXt-Small"),
    ]
    table1 = []
    for key, label in backbone_order:
        no = clahe[(clahe.backbone == key) & (clahe.clahe == False)].iloc[0]
        yes = clahe[(clahe.backbone == key) & (clahe.clahe == True)].iloc[0]
        delta = yes.acc_mean - no.acc_mean
        table1.append([label, fmt_pm(no.acc_mean, no.acc_std), fmt_pm(yes.acc_mean, yes.acc_std), f"{delta:.2f}"])

    tests = results["second"]["paired_seed_tests"]
    labels = {
        "A_gap": "A: Global average pooling",
        "B_raw_bilinear": "B: Mean-preserving second-order",
        "C_gap_raw_bilinear": "C: GAP + second-order fusion",
        "D_centered_cov": "D: Centered covariance",
    }
    table2 = []
    for arm in ["A_gap", "B_raw_bilinear", "C_gap_raw_bilinear", "D_centered_cov"]:
        row = summary[arm]
        if arm == "A_gap":
            table2.append([labels[arm], fmt_pm(row["class_acc_mean"], row["class_acc_std"]), "--", "--", f'{row["severe_classderived_mean"]:.1f}'])
        else:
            test = tests[arm]
            table2.append([
                labels[arm],
                fmt_pm(row["class_acc_mean"], row["class_acc_std"]),
                f'{test["mean_delta_pp"]:.2f}',
                f'{test["p_value"]:.3f}',
                f'{row["severe_classderived_mean"]:.1f}',
            ])

    ordinal = results["ordinal"]
    table3 = [
        ["Flat 24-class CE", "88.61", "0.908", "0.114", "6 (0.68%)", "--"],
        ["Hierarchical ordinal", "89.41", "0.917 (CORAL)", "0.108 (CORAL)", "2 (0.23%)", "+0.95 pp, p = 0.47"],
    ]
    table4 = [
        ["Prasetyo et al. ResNet50 + CLAHE [5]", "78.82", "single run", "Peer-reviewed FFE benchmark"],
        ["Yildiz et al. VGG19 + ANN [6]", "77.30", "single run", "Three-class FFE freshness model"],
        ["Hoang et al. Swin-Tiny + RF/LGBM [7]", "85.99", "single run", "Published prior best; different split"],
        ["This study: GAP Swin-Tiny", "88.53 +/- 0.75", "5 seeds", "Validated baseline, no CLAHE"],
    ]
    return {
        "table_model": table_model,
        "table0": table0,
        "table1": table1,
        "table2": table2,
        "table3": table3,
        "table4": table4,
    }


REFERENCES = [
    "S. Sattar, T. Abbas, M. Tabish, and G. Zhiqiang, \"Computer vision in aquaculture: transforming fish freshness monitoring,\" Critical Reviews in Food Science and Nutrition, pp. 1-29, 2025, doi: 10.1080/10408398.2025.2607533.",
    "P. K. Prabhakar, S. Vatsa, P. P. Srivastav, and S. S. Pathak, \"A comprehensive review on freshness of fish and assessment: analytical methods and recent innovations,\" Food Research International, vol. 133, p. 109157, 2020, doi: 10.1016/j.foodres.2020.109157.",
    "T. Murakoshi, T. Masuda, K. Utsumi, K. Tsubota, and Y. Wada, \"Glossiness and perishable food quality: visual freshness judgment of fish eyes based on luminance distribution,\" PLoS ONE, vol. 8, no. 3, p. e58994, 2013, doi: 10.1371/journal.pone.0058994.",
    "E. Prasetyo, R. D. Adityo, N. Suciati, and C. Fatichah, \"The Freshness of the Fish Eyes Dataset,\" Mendeley Data, 2020, doi: 10.17632/xzyx7pbr3w.1.",
    "E. Prasetyo, R. Purbaningtyas, R. D. Adityo, N. Suciati, and C. Fatichah, \"Combining MobileNetV1 and depthwise separable convolution bottleneck with Expansion for classifying the freshness of fish eyes,\" Information Processing in Agriculture, vol. 9, no. 4, pp. 485-496, 2022, doi: 10.1016/j.inpa.2022.01.002.",
    "M. B. Yildiz, E. T. Yasin, and M. Koklu, \"Fisheye freshness detection using common deep learning algorithms and machine learning methods with a developed mobile application,\" European Food Research and Technology, vol. 250, pp. 1919-1932, 2024, doi: 10.1007/s00217-024-04493-0.",
    "P.-H. Hoang, N.-T. Trinh, V.-M. Tran, and T.-T.-H. Phan, \"Deep feature optimization for enhanced fish freshness assessment,\" Ecological Informatics, vol. 95, p. 103711, 2026, doi: 10.1016/j.ecoinf.2026.103711.",
    "A. Banwari, R. C. Joshi, N. Sengar, and M. K. Dutta, \"Computer vision technique for freshness estimation from segmented eye of fish image,\" Ecological Informatics, vol. 69, p. 101602, 2022, doi: 10.1016/j.ecoinf.2022.101602.",
    "H. Mohammadi Lalabadi, M. Sadeghi, and S. A. Mireei, \"Fish freshness categorization from eyes and gills color features using multi-class artificial neural network and support vector machines,\" Aquacultural Engineering, vol. 90, p. 102076, 2020, doi: 10.1016/j.aquaeng.2020.102076.",
    "K. Zuiderveld, \"Contrast limited adaptive histogram equalization,\" in Graphics Gems IV, P. S. Heckbert, Ed. Academic Press, 1994, pp. 474-485.",
    "K. He, X. Zhang, S. Ren, and J. Sun, \"Deep residual learning for image recognition,\" in Proc. IEEE CVPR, 2016, pp. 770-778, doi: 10.1109/CVPR.2016.90.",
    "Z. Liu, H. Mao, C.-Y. Wu, C. Feichtenhofer, T. Darrell, and S. Xie, \"A ConvNet for the 2020s,\" in Proc. IEEE/CVF CVPR, 2022, pp. 11976-11986, doi: 10.1109/CVPR52688.2022.01167.",
    "Z. Liu et al., \"Swin Transformer: hierarchical vision transformer using shifted windows,\" in Proc. IEEE/CVF ICCV, 2021, pp. 10012-10022, doi: 10.1109/ICCV48922.2021.00986.",
    "T.-Y. Lin, A. RoyChowdhury, and S. Maji, \"Bilinear CNN models for fine-grained visual recognition,\" in Proc. IEEE ICCV, 2015, pp. 1449-1457, doi: 10.1109/ICCV.2015.170.",
    "P. Li, J. Xie, Q. Wang, and Z. Gao, \"Towards faster training of global covariance pooling networks by iterative matrix square root normalization,\" in Proc. IEEE/CVF CVPR, 2018, pp. 947-955, doi: 10.1109/CVPR.2018.00105.",
    "W. Cao, V. Mirjalili, and S. Raschka, \"Rank consistent ordinal regression for neural networks with application to age estimation,\" Pattern Recognition Letters, vol. 140, pp. 325-331, 2020, doi: 10.1016/j.patrec.2020.11.008.",
    "P. Geurts, D. Ernst, and L. Wehenkel, \"Extremely randomized trees,\" Machine Learning, vol. 63, no. 1, pp. 3-42, 2006, doi: 10.1007/s10994-006-6226-1.",
    "R. R. Selvaraju et al., \"Grad-CAM: visual explanations from deep networks via gradient-based localization,\" in Proc. IEEE ICCV, 2017, pp. 618-626, doi: 10.1109/ICCV.2017.74.",
    "R. Wightman, \"PyTorch Image Models,\" GitHub repository, 2019. [Online]. Available: https://github.com/huggingface/pytorch-image-models.",
]


def add_references(doc: Document) -> None:
    add_h1(doc, "REFERENCES")
    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"[{i}] {ref}")
        run.font.name = FONT
        run.font.size = Pt(8)


def build_doc() -> None:
    run_figure_builder()
    results = load_results()
    tables = build_tables(results)

    doc = Document()
    apply_document_style(doc)

    add_title(doc, TITLE_ID)
    add_title(doc, TITLE_EN)
    add_center(doc, "Dason Tiovino#1, [Supervisor Name]*2", size=11, bold=True)
    add_center(doc, "#Informatics Study Program, Pradita University", size=10, italic=True)
    add_center(doc, "Tangerang, Banten, Indonesia", size=10, italic=True)
    add_center(doc, "1d4sontiovino@gmail.com", size=9, italic=True)
    add_center(doc, "*[Supervisor Department], [Supervisor University]", size=10, italic=True)
    add_center(doc, "2[supervisor-email@example.com]", size=9, italic=True)
    add_center(doc, "Corresponding author: d4sontiovino@gmail.com", size=9)

    add_h2(doc, "Abstrak")
    add_paragraph(
        doc,
        "Penelitian ini memvalidasi Swin-Tiny dengan global average pooling (GAP) dan head 24-kelas sebagai model utama untuk klasifikasi kesegaran mata ikan pada dataset Freshness of Fish Eyes (FFE). "
        "Model akhir dilatih tanpa CLAHE karena hipotesis penelitian menyatakan bahwa kesegaran FFE terutama dibawa oleh statistik okular global, bukan kontras lokal. "
        "Untuk menguji hipotesis tersebut, penelitian ini mengukur statistik luminansi dan warna pusat mata, lalu menjalankan ablasi CLAHE, pooling orde-dua, dan struktur label ordinal. "
        "Model GAP Swin-Tiny mencapai akurasi 24-kelas sebesar 88,53 +/- 0,75% pada protokol lima seed, lebih tinggi 9,71 poin daripada benchmark ResNet50 terpublikasi dan 2,54 poin daripada angka Hoang et al.; perbandingan Hoang diperlakukan sebagai indikatif karena protokol split tidak identik. "
        "Analisis statistik global menunjukkan bahwa fitur `raw_center_lab_l_std` berkorelasi dengan peringkat kesegaran pada semua delapan spesies (rho = 0,245, p < 1e-60), dan fitur global mentah mengklasifikasikan kesegaran sebesar 58,13 +/- 0,87%, jauh di atas baseline spesies saja sebesar 37,06%. "
        "CLAHE tidak meningkatkan backbone mana pun, menurunkan ResNet50 sebesar 3,34 poin, dan juga melemahkan statistik global mentah sebesar 2,35 poin (p = 0,005). "
        "Semua varian pooling orde-dua lebih buruk daripada GAP; kandidat terbaik turun 1,68 poin (p = 0,007). "
        "Model ordinal hierarkis tidak meningkatkan akurasi secara signifikan, tetapi memberi pembacaan kesegaran terkalibrasi dengan QWK 0,917. "
        "Hasil ini mendukung rekomendasi praktis: gunakan GAP Swin-Tiny tanpa CLAHE untuk FFE.",
        size=9,
    )
    add_paragraph(doc, "Kata kunci-- CLAHE; GAP Swin-Tiny; global ocular statistics; klasifikasi kesegaran ikan; second-order pooling", indent=False, size=9, italic=True)

    add_h2(doc, "Abstract")
    add_paragraph(
        doc,
        "This study validates Swin-Tiny with global average pooling (GAP) and a 24-class head as the main model for fish-eye freshness classification on the Freshness of Fish Eyes (FFE) dataset. "
        "The final model is trained without CLAHE because the central hypothesis is that FFE freshness is carried mainly by global ocular statistics, not local contrast. "
        "To test this hypothesis, we measure central-eye luminance and color statistics, then run controlled ablations on CLAHE, second-order pooling, and ordinal label structure. "
        "The GAP Swin-Tiny model reaches 88.53 +/- 0.75% 24-class accuracy under a five-seed protocol, 9.71 points above the published ResNet50 benchmark and 2.54 points above the value reported by Hoang et al.; the Hoang comparison is treated as indicative because the split protocol is not identical. "
        "The global-statistics analysis shows that `raw_center_lab_l_std` correlates with freshness rank in all eight species (rho = 0.245, p < 1e-60), and raw global features classify freshness at 58.13 +/- 0.87%, far above a species-only baseline of 37.06%. "
        "CLAHE improves no backbone, reduces ResNet50 by 3.34 points, and weakens raw global statistics by 2.35 points (p = 0.005). "
        "Every second-order pooling variant performs worse than GAP; the best candidate drops 1.68 points (p = 0.007). "
        "Hierarchical ordinal learning does not significantly improve accuracy, but it provides a calibrated freshness readout with QWK 0.917. "
        "These results support a practical recommendation: use GAP Swin-Tiny without CLAHE for FFE.",
        size=9,
    )
    add_paragraph(doc, "Keywords-- CLAHE; fish freshness classification; GAP Swin-Tiny; global ocular statistics; second-order pooling", indent=False, size=9, italic=True)

    add_h1(doc, "I. INTRODUCTION")
    add_paragraph(doc, "Fish freshness grading matters because seafood quality declines quickly after harvest, and manual sensory inspection is subjective, slow, and hard to standardize across markets. Reviews of fish freshness assessment describe the need for rapid and non-destructive methods that can support quality control along the supply chain [1], [2]. The fish eye is a useful visual indicator because freshness loss changes corneal transparency, gloss, and iris color. Murakoshi et al. showed that perceived fish-eye freshness is related to luminance distribution, which directly supports a global optical interpretation of the signal [3].")
    add_paragraph(doc, "The public Freshness of Fish Eyes (FFE) dataset enables controlled study of this problem. It contains 4,390 usable eye images in this repository, organized as eight species and three ordered freshness levels, giving 24 species-freshness classes [4]. Prasetyo et al. established the main peer-reviewed benchmark using CNN models on FFE; their ResNet50 comparison reached 78.82%, while their proposed MB-BE model reached 63.21% [5]. Yildiz et al. later evaluated deep feature extraction and classical machine-learning classifiers for three-level FFE freshness classification [6]. Hoang et al. published a stronger deep-feature optimization framework in 2026, reporting 85.99% using Swin-Tiny features with Random Forest and LightGBM-based feature selection [7].")
    add_paragraph(doc, "These prior results improved accuracy, but they did not isolate which representation should be used as the final practical model for FFE. This gap matters. A stronger model claim should explain why the chosen architecture and preprocessing are appropriate, not only report a higher number. This study therefore validates a simple final classifier--GAP Swin-Tiny with a 24-class species-freshness head and no CLAHE--and then tests the visual-signal assumptions behind that choice.")
    add_paragraph(doc, "The central hypothesis is simple: FFE freshness is dominated by global ocular statistics. If that hypothesis is correct, the useful signal should be measurable with simple ocular luminance and color statistics, GAP should be a suitable aggregation mechanism, CLAHE should not help, and second-order texture pooling should not outperform GAP. If the hypothesis is wrong, one of those alternatives should provide a clear improvement.")
    add_paragraph(doc, "The experiments are therefore ordered as a model-validation chain. First, the final GAP Swin-Tiny classifier is defined and evaluated under a five-seed protocol. Second, direct global-statistics analysis checks whether the signal behind the model is measurable without deep learning. Third, CLAHE tests whether local contrast helps or damages that signal. Fourth, second-order pooling tests whether higher-order texture adds value beyond GAP. Fifth, hierarchical ordinal learning tests whether label structure improves accuracy or only improves readout. This order makes the model the endpoint and the ablations the defense.")
    add_paragraph(doc, "The study also raises the statistical standard for this dataset. Several FFE papers report one final accuracy value, but a single run does not reveal seed variance, paired prediction differences, or whether a small improvement is meaningful. This work reports mean and standard deviation across seeds where possible, uses paired t-tests for shared-seed model comparisons, and uses McNemar's exact test when two classifiers are evaluated on the same test samples. For ordinal freshness, it reports QWK, MAE, and severe-error counts because a Not Fresh versus Highly Fresh confusion has a different practical cost from an adjacent Fresh versus Highly Fresh confusion.")
    add_paragraph(doc, "The results support this model choice. GAP Swin-Tiny reaches 88.53 +/- 0.75% 24-class accuracy across five seeds. Direct global statistics show that broad central-eye luminance features carry freshness information. CLAHE gives no benefit and harms ResNet50. Second-order pooling significantly underperforms GAP. Hierarchical ordinal supervision does not significantly improve 24-class accuracy, although it adds a calibrated freshness readout. The contribution is therefore a validated final model plus the evidence explaining why the simpler GAP design is preferable to the tested add-ons.")

    add_h1(doc, "II. METHODS")
    add_h2(doc, "A. Dataset and Splits")
    add_paragraph(doc, "All experiments use the FFE dataset [4]. The class folders follow the format Species - Freshness Level. The species are Chanos Chanos, Johnius Trachycephalus, Nibea Albiflora, Rastrelliger Faughni, Upeneus Moluccensis, Eleutheronema Tetradactylum, Oreochromis Mossambicus, and Oreochromis Niloticus. The freshness levels are Highly Fresh, Fresh, and Not Fresh. Images are resized to 224 x 224 pixels and normalized with ImageNet statistics.")
    add_figure(doc, "results/figures/jutisi/figure1_ffe_samples.png", "Figure 1. Representative FFE eye images across freshness levels.", 5.7)
    add_paragraph(doc, "The CLAHE and label-structure studies use fixed 60/20/20 train-validation-test splits unless stated otherwise. The second-order pooling study uses a true 64/16/20 split with five shared seeds. These protocols are reported separately because cross-protocol accuracy values are not strict head-to-head comparisons.")

    add_h2(doc, "B. Final Model and Common Training Recipe")
    add_paragraph(doc, FINAL_MODEL_STATEMENT + " The 24 classes are the Cartesian product of eight species and three freshness levels, so the classifier learns species-freshness categories directly. Freshness-level metrics are then derived from the predicted 24-class label. This design keeps benchmark comparability while still allowing ordinal freshness analysis.")
    add_paragraph(doc, "Models use ImageNet-pretrained backbones from timm [19], AdamW with learning rate 2e-4 and weight decay 0.05, a five-epoch linear warmup followed by cosine decay, label smoothing 0.1, dropout 0.3, stochastic depth 0.1, mixed precision, batch size 64, and early stopping with patience 18 up to 90 epochs. Augmentation uses horizontal flip, mild rotation, and brightness/contrast jitter. Hue and saturation shifts are avoided because color and opacity are part of the freshness signal.")
    add_paragraph(doc, "The recipe was kept intentionally conservative. The goal was not to search every possible training trick, but to remove weak training as an explanation for the ablation results. Full fine-tuning allows the pretrained backbone to adapt to FFE. Warmup prevents unstable early updates. Cosine decay reduces the learning rate smoothly without manually selecting late-stage schedules. Label smoothing reduces overconfidence on 24 fine-grained classes with limited samples per class. These choices are standard enough to support a fair comparison, while still avoiding transformations that would alter the biological freshness cues.")
    add_paragraph(doc, "All reported comparisons change one design axis at a time. In the preprocessing study, the paired factor is CLAHE. In the pooling study, the paired factor is the feature-pooling operator. In the label-structure study, the paired factor is whether auxiliary species and ordinal freshness losses are active. This is stricter than simply comparing final systems because it makes the interpretation of each delta clearer.")

    add_h2(doc, "C. Direct Global-Signal Analysis")
    add_paragraph(doc, "Before interpreting CNN ablations, we quantify the visual signal directly. For each FFE image, we extract scalar statistics from raw full images, raw central-eye crops, and CLAHE-processed central-eye crops. The central crop uses 65% of image width and height as a lightweight ocular-region proxy; it is not a full cornea segmentation. For each region, we compute grayscale, CIE Lab, and HSV statistics: mean, standard deviation, skewness, selected percentiles, percentile range, and highlight contrast. This follows the luminance-distribution motivation of Murakoshi et al. [3], but applies it directly to the FFE benchmark and ties it to model ablations.")
    add_paragraph(doc, "We evaluate these statistics in two ways. First, Spearman correlations test whether each feature tracks ordered freshness rank globally and within each species. Species-stratified consistency is important because species color differences can otherwise create false correlations. Second, a balanced logistic-regression classifier predicts the three freshness levels from global statistics over five stratified 80/20 splits. We compare raw full-image statistics, raw central-crop statistics, CLAHE central-crop statistics, species-only labels, and species plus global statistics. This analysis is intentionally simple: if global ocular statistics carry no direct signal, the paper's main interpretation should not be claimed.")

    add_h2(doc, "D. Experiment A: Preprocessing Axis")
    add_paragraph(doc, "Experiment A tests whether CLAHE helps FFE. We compare CLAHE and no-CLAHE on ResNet50 [11], EfficientNetV2-S, and ConvNeXt-Small [12] across three seeds, for 18 runs. CLAHE is applied to the luminance channel in CIE Lab with clip limit 2.0 and tile grid 8 x 8. All other factors are paired. The allowed conclusion is per-backbone change in mean test accuracy, not a new final model claim.")

    add_h2(doc, "E. Experiment B: Feature-Pooling Axis")
    add_paragraph(doc, "Experiment B tests whether second-order texture statistics improve Swin-Tiny [13]. The baseline arm uses GAP. The first candidate uses an uncentered bilinear Gram matrix after a 1 x 1 projection to 128 channels, followed by signed square-root and L2 normalization. The second candidate concatenates GAP and the uncentered second-order vector. The diagnostic control uses centered covariance, which removes the channel mean before computing the Gram matrix. Each arm uses five seeds: 42, 123, 2024, 7, and 2025.")
    add_paragraph(doc, "The predefined success rule requires a candidate to beat GAP by at least 1.0 percentage point and pass a paired t-test with p < 0.05. A non-positive mean delta is a kill condition. This rule prevents treating a failed add-on as a positive method.")
    add_paragraph(doc, "The centered covariance control is the key diagnostic part of the design. If freshness were mainly texture, removing the mean should not be fatal because covariance would still capture co-occurrence structure. If freshness were mainly global opacity or color level, removing the mean should hurt. Therefore, the relative ordering among uncentered bilinear pooling, GAP-plus-bilinear fusion, and centered covariance gives more information than accuracy alone.")

    add_h2(doc, "F. Experiment C: Label-Structure Axis")
    add_paragraph(doc, "Experiment C tests whether encoding the biological label structure improves the classifier. FFE labels factor as eight species and three ordered freshness levels. The hierarchical ordinal model keeps the 24-class head for benchmark comparability, adds an eight-class species head, and adds a two-threshold CORAL freshness head for the ordered labels Not Fresh < Fresh < Highly Fresh [16]. The total loss is L = Lclass + 0.3 Lspecies + 0.7 Lcoral.")
    add_paragraph(doc, "The flat control uses the same Swin-Tiny backbone and training recipe, but species and freshness losses are set to zero. Freshness metrics for both models are derived from the 24-class argmax to avoid comparing a trained CORAL head against an untrained auxiliary head. We report 24-class accuracy, quadratic weighted kappa (QWK), mean absolute error (MAE), and severe errors, where severe means a rank distance of two.")

    add_h2(doc, "G. Explainability")
    add_paragraph(doc, "Grad-CAM is used as qualitative xAI support [18]. It does not prove causality, but it checks whether representative trained arms attend to the ocular region rather than unrelated background. The xAI panel uses the same three representative freshness samples across three columns: the CLAHE model, the primary second-order pooling candidate, and the final GAP Swin-Tiny model.")
    add_paragraph(doc, "The xAI analysis is deliberately secondary. A heatmap cannot replace controlled accuracy and error analysis. Its role is to make the model behavior auditable: if the rejected add-ons or the final model consistently attended to labels, borders, or background, the numerical results would be scientifically weaker. In this manuscript, Grad-CAM supports the interpretation that the trained models rely on the eye region, while the ablation tables provide the actual evidence.")

    add_h1(doc, "III. RESULTS AND DISCUSSION")
    add_paragraph(doc, "The results are read as one connected argument. First, the final GAP Swin-Tiny model is stated and evaluated. Second, direct global-statistics analysis explains why GAP is a suitable design. Third, the CLAHE and second-order experiments reject alternative assumptions. Fourth, the ordinal and xAI analyses show how the result can be read and audited. This order keeps the paper from becoming a list of unrelated experiments.")
    add_h2(doc, "A. GAP Swin-Tiny Is the Final Recommended Model")
    add_paragraph(doc, "Table 1 and Figure 2 show the endpoint of the study. The recommended model is not CLAHE, not second-order pooling, and not a hybrid feature-selection stack. It is a plain GAP Swin-Tiny classifier trained end-to-end on the 24 species-freshness classes without CLAHE.")
    add_table(doc, "Table 1. Final model specification and validated performance.", ["Component", "Choice or Result", "Role in Paper"], tables["table_model"])
    add_figure(doc, "results/figures/jutisi/figure4_final_model_context.png", "Figure 2. Final GAP Swin-Tiny result in FFE prior-work context.", 5.8)
    add_paragraph(doc, "The five-seed result is 88.53 +/- 0.75% 24-class accuracy. This is 9.71 points above the peer-reviewed ResNet50 benchmark of Prasetyo et al. and 2.54 points above the 85.99% reported by Hoang et al. The Hoang comparison is not written as a definitive head-to-head victory because the exact split and protocol are not reproduced. The defensible claim is narrower and stronger: this work reports a higher multi-seed validated GAP Swin-Tiny baseline and then explains why the simpler GAP design is appropriate.")

    add_h2(doc, "B. Direct Global Statistics Explain the Model Choice")
    add_paragraph(doc, "Table 2 and Figures 3-4 show the direct signal analysis. The strongest central-eye luminance statistic is Lab-L standard deviation: it correlates with ordered freshness rank (rho = 0.245, p < 1e-60) and has the same direction in all eight species. Other high-ranked features are also luminance-distribution ranges, including Lab-L p90-p10 and gray-level standard deviation. This matters because it rules out a purely species-driven explanation: the signal remains visible inside each species. The correlation magnitude itself is modest -- rho near 0.25 explains only a small share of variance -- so the interpretation deliberately does not rest on the very small p-value, which mainly reflects the large sample size. The decisive evidence is the eight-species sign consistency, which a species-confound account cannot produce, together with the feature-only classifier gain reported next.")
    add_table(doc, "Table 2. Direct global ocular signal analysis on FFE.", ["Evidence", "Result", "Interpretation"], tables["table0"])
    add_figure(doc, "results/figures/jutisi/figure2_global_signal_distributions.png", "Figure 3. Central-eye global statistics across ordered freshness levels.", 5.8)
    add_figure(doc, "results/figures/jutisi/figure3_global_signal_classifier.png", "Figure 4. Feature-only freshness classifiers using global ocular statistics.", 5.8)
    add_paragraph(doc, "The feature-only classifier gives the practical check. Species-only information reaches only 37.06% freshness accuracy, while raw central global statistics reach 58.13 +/- 0.87% with QWK 0.418. This is not strong enough to replace the CNN, but it is strong enough to prove that simple global ocular statistics carry real freshness information. CLAHE central statistics drop to 55.79 +/- 1.59%, and raw statistics beat CLAHE statistics by 2.35 points with a paired p-value of 0.005. Therefore, the later CLAHE failure is not just a CNN artifact; CLAHE weakens the same global signal measured directly.")

    add_h2(doc, "C. CLAHE Does Not Help FFE")
    add_paragraph(doc, "Table 3 and Figure 5 show the CLAHE ablation. CLAHE improves no backbone. ResNet50 drops by 3.34 percentage points, while EfficientNetV2-S is unchanged within noise and ConvNeXt-Small drops by 0.98 points. The pattern now has direct support from Table 2: local tile-wise equalization is mismatched to a signal carried by global opacity and luminance distribution.")
    add_table(doc, "Table 3. CLAHE ablation on FFE (test accuracy, mean +/- SD over three seeds).", ["Backbone", "No CLAHE (%)", "With CLAHE (%)", "Delta (pp)"], tables["table1"])
    add_figure(doc, "results/figures/jutisi/figure2_clahe_ablation.png", "Figure 5. CLAHE versus no-CLAHE accuracy by backbone.", 5.8)
    add_paragraph(doc, "This result is important for publication because it directly corrects a common assumption in the benchmark line. CLAHE is not a harmless default. It can erase the luminance differences that represent eye turbidity. For FFE-style data, color-safe augmentation is a safer default than local histogram equalization.")
    add_paragraph(doc, "The result also explains why the ResNet50 degradation is the strongest. ResNet50 uses conventional convolutional stages that learn many early contrast and edge responses. When CLAHE changes local luminance statistics before training, it changes the input distribution of those early filters. ConvNeXt-Small is less damaged because its larger depthwise kernels and modern normalization can integrate wider context, but even this backbone does not benefit. The conclusion is therefore not architecture-specific: no tested backbone gains from local histogram equalization.")
    add_paragraph(doc, "One scope point deserves explicit statement. The final recommended model is Swin-Tiny, but the paired CLAHE ablation in Table 3 was run on ResNet50, EfficientNetV2-S, and ConvNeXt-Small rather than on Swin-Tiny itself, so the no-CLAHE decision for the final model is not justified by a within-Swin ablation. It rests instead on two model-independent lines of evidence. First, the direct global-statistics analysis in Table 2 shows that CLAHE weakens the freshness signal before any network is trained, reducing raw central-statistic freshness accuracy by 2.35 points (p = 0.005); this is a property of the input transform, not of a particular backbone. Second, CLAHE produced no accuracy gain on any of three architecturally diverse backbones spanning residual-convolutional, mobile-inverted-residual, and modern large-kernel designs. Because the harm is visible at the signal level and the no-benefit pattern is architecture-independent, applying CLAHE to Swin-Tiny would be expected to follow the same pattern; a within-Swin CLAHE pair remains a useful confirmatory run for future work.")
    add_paragraph(doc, "This first result sets up the second experiment. If local contrast is not the missing signal, the next plausible explanation is that freshness may require higher-order texture or feature co-occurrence beyond the mean activation. The pooling study tests exactly that alternative while keeping the same broad visual interpretation in view.")

    add_h2(doc, "D. GAP Beats Second-Order Pooling")
    add_paragraph(doc, "Table 4 and Figure 6 show the pooling study. GAP reaches 88.53 +/- 0.75%. All second-order variants are significantly worse. The mean-preserving bilinear variant drops by 1.68 points with p = 0.007. The GAP-plus-second-order fusion drops by 2.60 points with p = 0.037. The centered covariance control drops by 2.73 points with p = 0.028.")
    add_table(doc, "Table 4. Swin-Tiny pooling comparison (five seeds, 64/16/20 split).", ["Pooling Operator", "Accuracy (%)", "Delta vs GAP", "p", "Severe Errors"], tables["table2"])
    add_figure(doc, "results/figures/jutisi/figure3_pooling_comparison.png", "Figure 6. Pooling-operator comparison with per-seed points and mean accuracy.", 5.8)
    add_paragraph(doc, "The internal pattern is also informative. The centered covariance control is the worst second-order arm, while the uncentered mean-preserving variant is least harmful. This means that removing the mean damages the representation more than adding texture helps it. The result refutes the higher-order texture hypothesis for this benchmark and supports GAP as the correct inductive bias.")
    add_paragraph(doc, "This negative result is stronger than a simple failure to improve. The uncentered second-order arm was designed to preserve the mean while adding texture. It was the most favorable candidate for the mean-plus-texture hypothesis, yet it still lost significantly. The fusion arm also lost, which means the model did not merely need GAP and texture together. The likely explanation is redundancy and overfitting: Swin attention already models interactions among tokens, while the Gram vector adds thousands of features on a small training set. The practical recommendation is to keep the head simple unless a larger or more diverse dataset justifies the added capacity.")
    add_paragraph(doc, "This second result strengthens the same story as the CLAHE result. CLAHE fails because local contrast is not the right cue. Second-order pooling fails because higher-order texture is also not the right cue. Both failures point back to the same mechanism: the discriminative information is carried mainly by global mean-level properties of the eye representation.")

    add_h2(doc, "E. Ordinal Supervision Adds Readout, Not Accuracy")
    add_paragraph(doc, "This final ablation is a secondary, capability-oriented check rather than a core part of the representational argument; it asks only whether label structure changes the conclusion drawn from the preprocessing and pooling experiments. Table 5 shows the label-structure ablation. The hierarchical ordinal model improves the best-checkpoint 24-class accuracy from 88.61% to 89.41%, but the three-seed paired effect is not significant (p = 0.47), and McNemar's exact test on the best checkpoints is also not significant (p = 0.44). Severe errors fall from 6 to 2, but that reduction is also not significant (p = 0.29). The 88.61% flat figure here is a best-checkpoint accuracy from the three-seed 60/20/20 label-structure study and is therefore not identical to the 88.53 +/- 0.75% five-seed 64/16/20 headline in Table 1; the two values describe the same backbone under different protocols and are reported separately rather than merged.")
    add_table(doc, "Table 5. Label-structure ablation on the 878-image test set.", ["Model", "24-Class Acc. (%)", "Freshness QWK", "Freshness MAE", "Severe Errors", "Accuracy Test"], tables["table3"])
    add_paragraph(doc, "This is not a failure if the claim is stated correctly. The ordinal model should not be sold as an accuracy winner. Its defensible value is a calibrated, rank-monotonic freshness readout: the CORAL head reaches QWK 0.917, MAE 0.108, and a severe-error rate of 0.23%. That readout is useful for quality grading because it makes freshness explicit rather than hidden inside a 24-class label.")
    add_paragraph(doc, "The severe-error metric is especially relevant for fish freshness. A model that confuses Fresh and Highly Fresh is imperfect, but the practical risk is lower than predicting Highly Fresh for a Not Fresh sample. Accuracy treats both mistakes as one error. QWK and MAE recover the ordering information, and severe-error counts expose the safety-critical cases. This is why the ordinal readout remains useful even without a statistically significant accuracy gain.")
    add_paragraph(doc, "The ordinal experiment therefore completes the chain rather than contradicting it. It shows that the global-first-order representation is already strong for 24-class accuracy, while structured supervision is better understood as an output calibration layer. In other words, label structure improves how the result can be read by humans, but it does not overturn the representational conclusion from the first two experiments.")
    add_figure(doc, "results/figures/hierarchical_ordinal/hierarchical_ordinal_freshness_confusion_matrix.png", "Figure 7. Freshness confusion matrix for the hierarchical ordinal readout.", 4.8)

    add_h2(doc, "F. Comparison with Prior FFE Studies")
    add_paragraph(doc, "Table 6 places the validated GAP Swin-Tiny baseline against prior FFE studies. The 88.53 +/- 0.75% result is 9.71 points above the ResNet50 benchmark reported by Prasetyo et al. [5] and 2.54 points above the published 85.99% result of Hoang et al. [7]. The Hoang comparison is informative but not a strict head-to-head claim because the split details are not identical. The stronger claim is that this work reports a multi-seed validated baseline with paired ablations, which prior FFE papers generally do not provide.")
    add_table(doc, "Table 6. FFE results in context.", ["Method", "Accuracy (%)", "Seeds", "Status"], tables["table4"])
    add_paragraph(doc, "This distinction matters for reviewer trust. A manuscript can honestly say that its validated baseline is higher than the reported prior number, but it should not use that fact to claim a final state-of-the-art replacement unless the test partition, split seed, and evaluation protocol are identical. The contribution here is the controlled evidence around the representation, not an overstated leaderboard claim.")

    add_h2(doc, "G. Stage-Wise Grad-CAM: Correct vs Incorrect Predictions")
    add_paragraph(doc, "Figures 8 and 9 audit the full evidence chain visually with Grad-CAM. Both panels share the same three columns -- the CLAHE EfficientNetV2-S model, the second-order Swin-Tiny candidate, and the final GAP Swin-Tiny model -- and the same three freshness rows, so the preprocessing, pooling, and final-model stages are read side by side under one figure logic. To make the comparison sharper than a single mixed grid, correctly classified and misclassified cases are separated: Figure 8 shows one correctly classified example per cell, and Figure 9 shows one misclassified example per cell. These are representative qualitative cases selected for attention auditing, not a quantitative test partition; the tables carry the statistical evidence.")
    add_figure(doc, "results/figures/jutisi/figure8_xai_correct_3x3.png", "Figure 8. Correctly classified samples: Grad-CAM for the CLAHE, second-order, and final GAP Swin-Tiny stages across the three freshness levels.", 5.6)
    add_paragraph(doc, "On the correctly classified samples in Figure 8, the three stages attend in visibly different ways that match their measured behavior, which is why the three experiments are read as one correlated chain rather than as separate tests. The CLAHE model concentrates on a single saturated high-contrast spot of the eye, consistent with a representation pulled toward local contrast rather than the broad ocular region. The second-order candidate produces sparse, fragmented attention, consistent with a Gram representation that emphasizes isolated texture co-occurrences. The final GAP Swin-Tiny model spreads attention across the cornea and iris as one connected region, which is the behavior expected if freshness is carried by global ocular statistics. The qualitative pattern therefore lines up with the quantitative results: CLAHE (preprocessing axis) and second-order pooling (feature axis) both move attention away from the global ocular region, while GAP keeps it there, and the global-statistics analysis in Section III-B already showed that this same global region carries the freshness signal.")
    add_figure(doc, "results/figures/jutisi/figure9_xai_incorrect_3x3.png", "Figure 9. Misclassified samples: Grad-CAM for the same three stages, annotated with the true and predicted freshness level.", 5.6)
    add_paragraph(doc, "Figure 9 shows where each stage fails and ties the visual audit back to the safety-relevant metrics. The misclassifications include several severe Not Fresh versus Highly Fresh flips, with confident errors above 0.85 softmax probability for both the CLAHE model and the final GAP model. Two points follow. First, even the final model can make a high-confidence severe error, which is exactly why the ordinal readout and severe-error metric in Section III-E matter for deployment beyond raw accuracy. Second, on the failures the CLAHE and second-order stages keep their narrow or fragmented attention, while the GAP model fails with eye-centered but over-spread attention rather than by leaving the eye; this is consistent with the global-first-order account rather than against it. As with Figure 8, these panels are interpretive support, and the controlled tables remain the primary evidence.")

    add_h2(doc, "H. Research Interpretation")
    add_paragraph(doc, "The experiments are connected by one logic. GAP Swin-Tiny is the final model. Direct global statistics explain why GAP is plausible by showing that luminance-distribution features carry freshness information. CLAHE then tests local contrast and fails, including on those same statistics. Second-order pooling tests higher-order texture and fails. Hierarchical ordinal learning tests label-structure supervision and does not significantly improve accuracy, although it adds a useful calibrated readout. The only mechanism that consistently holds is the global aggregation represented by GAP.")
    add_paragraph(doc, "The global-first-order interpretation is also consistent with the biology of fish-eye freshness. Spoilage makes the cornea less transparent and reduces the gloss of the eye surface. These are not isolated high-frequency textures; they are broad changes in brightness distribution, opacity, and color. A GAP-pooled representation is well matched to this kind of signal because it summarizes feature presence over the whole ocular region. CLAHE and centered covariance move the model away from that signal: CLAHE distorts local luminance, and centered covariance explicitly removes the mean.")
    add_paragraph(doc, "This does not mean future work is impossible. It means that further accuracy gains are unlikely to come from generic contrast enhancement or larger texture statistics alone. More promising directions are direct corneal or iris region isolation, better label quality, larger cross-dataset validation, and direct measurement of eye luminance distribution to test the global-first-order account more explicitly.")
    add_paragraph(doc, "The main limitation is that the direct global-statistics analysis uses a central-eye crop rather than a manually segmented cornea and iris. A stronger future study should segment those structures, measure luminance and saturation distributions inside them, and test whether those measurements correlate with the model logits and freshness labels. Another limitation is protocol heterogeneity: the CLAHE and ordinal studies use 60/20/20, while the pooling study uses 64/16/20. The manuscript therefore keeps each table within its own protocol and does not merge all scores into a single direct ranking.")

    add_h1(doc, "IV. CONCLUSION")
    add_paragraph(doc, "This study validated GAP Swin-Tiny with a 24-class species-freshness head as the recommended FFE fish-eye freshness classifier. The model reaches 88.53 +/- 0.75% accuracy across five seeds without CLAHE. Direct global-statistics analysis showed that central-eye luminance distribution features correlate with freshness rank across all eight species and classify freshness substantially better than species labels alone, so the global-statistics model serves as diagnostic evidence for the final CNN. CLAHE weakened these global statistics and gave no CNN benefit, harming ResNet50 by 3.34 percentage points. Second-order pooling was significantly worse than GAP under a five-seed protocol, with the best candidate falling by 1.68 points (p = 0.007). Hierarchical ordinal learning did not significantly improve 24-class accuracy, but it provided a calibrated freshness readout with QWK 0.917. The evidence supports a practical conclusion: for FFE, use raw-image GAP Swin-Tiny and avoid CLAHE or higher-order pooling add-ons unless a stronger segmented-eye analysis justifies them.")

    add_references(doc)
    doc.save(OUT_DOCX)
    print(f"saved {OUT_DOCX}")


if __name__ == "__main__":
    build_doc()
